import os
import subprocess
import sys

# Instalar python-docx si no existe
try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Configurar márgenes
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Configurar fuente general
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11.5)

# Título
h1 = doc.add_heading('Optimización de la Inversión en el Sector Agrario', level=1)
h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in h1.runs:
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.size = Pt(14)

doc.add_paragraph()

# Caso
caso = doc.add_paragraph()
caso.add_run('Caso: ').bold = True
caso.add_run('Usted ha sido contratado por el Ministerio de Desarrollo Agrario para diseñar la ficha técnica de una encuesta regional. El objetivo es diagnosticar la situación de los productores de café en una zona de alta productividad que cuenta con 12,500 productores registrados. Los resultados determinarán la cuantía de los fondos de apoyo para el próximo año. Se le solicita calcular los tamaños de muestra necesarios para dos indicadores críticos, considerando que el presupuesto para el trabajo de campo es limitado y cada encuesta aplicada cuesta $25.00.')

# Tarea 1
h2 = doc.add_heading('Tarea 1: Estimación de la Tasa de Acceso al Crédito', level=2)
for run in h2.runs: run.font.name = 'Times New Roman'; run.font.color.rgb = RGBColor(0,0,0); run.font.size = Pt(12)

p_info = doc.add_paragraph('El Ministerio necesita saber qué proporción de productores tiene deudas con la banca formal para diseñar un programa de refinanciamiento.\n')
p_info.add_run('Información disponible: ').bold = True
p_info.add_run('Un estudio piloto indica que aproximadamente el 35% de los productores tiene crédito activo.\n')
p_info.add_run('Requerimiento: ').bold = True
p_info.add_run('Trabaje con un Nivel de Confianza del 95% y un Margen de Error del 4%.')

# Tarea A
q_a = doc.add_paragraph()
q_a.add_run('a) Calcule el tamaño de muestra necesario considerando que la población es finita (N = 12,500)').bold = True

doc.add_paragraph().add_run('Solución a):').bold = True
doc.add_paragraph('Población (N): 12,500\nProporción esperada (p): 0.35\nProbabilidad de fracaso (q): 0.65\nConfianza (NC): 95%, cuyo valor Z es 1.96\nError (e): 0.04 (4%)')

inst = doc.add_paragraph('Fórmulas base para Word (Cópialas y presiona Alt + = para convertirlas automáticamente):')
inst.runs[0].font.color.rgb = RGBColor(100, 100, 100)
inst.runs[0].italic = True

doc.add_paragraph('n = \\frac{N \\cdot Z^2 \\cdot p \\cdot q}{(N-1) \\cdot e^2 + Z^2 \\cdot p \\cdot q}')
doc.add_paragraph('n = \\frac{12500 \\cdot (1.96)^2 \\cdot 0.35 \\cdot 0.65}{12499 \\cdot (0.04)^2 + (1.96)^2 \\cdot 0.35 \\cdot 0.65}')
doc.add_paragraph('n = \\frac{10924.55}{19.9984 + 0.873964} = 523.40')

ans_a = doc.add_paragraph()
ans_a.add_run('Respuesta: ').bold = True
ans_a.add_run('El tamaño de muestra necesario es de 524 encuestas.')

# Tarea B
q_b = doc.add_paragraph()
q_b.add_run('b) ¿Cuál sería el costo total de la recolección de datos para este indicador?').bold = True

doc.add_paragraph().add_run('Solución b):').bold = True
doc.add_paragraph('Costo Total = 524 encuestas * $25.00 = $13,100.00')

ans_b = doc.add_paragraph()
ans_b.add_run('Respuesta: ').bold = True
ans_b.add_run('El costo total será de $13,100.00.')

# Tarea 2
h2_2 = doc.add_heading('Tarea 2: Estimación del Ingreso Promedio Mensual', level=2)
for run in h2_2.runs: run.font.name = 'Times New Roman'; run.font.color.rgb = RGBColor(0,0,0); run.font.size = Pt(12)

p_info2 = doc.add_paragraph('Para calcular el impacto del subsidio, se requiere conocer el ingreso neto promedio mensual por hectárea.\n')
p_info2.add_run('Información disponible: ').bold = True
p_info2.add_run('En el censo agropecuario anterior, la desviación estándar del ingreso en esta zona fue de $180\n')
p_info2.add_run('Requerimiento: ').bold = True
p_info2.add_run('El Ministro solicita una precisión alta: Nivel de Confianza del 99% y un Error Máximo admisible de $15.')

# Tarea C
q_c = doc.add_paragraph()
q_c.add_run('c) Calcule el tamaño de muestra requerido.').bold = True

doc.add_paragraph().add_run('Solución c):').bold = True
doc.add_paragraph('Población (N): 12,500\nDesviación Estándar (\u03C3): $180\nConfianza (NC): 99%, cuyo valor Z es 2.576\nError Máximo (E): $15')

inst = doc.add_paragraph('Fórmulas base para Word:')
inst.runs[0].font.color.rgb = RGBColor(100, 100, 100); inst.runs[0].italic = True

doc.add_paragraph('n = \\frac{N \\cdot Z^2 \\cdot \\sigma^2}{(N-1) \\cdot E^2 + Z^2 \\cdot \\sigma^2}')
doc.add_paragraph('n = \\frac{12500 \\cdot (2.576)^2 \\cdot (180)^2}{12499 \\cdot (15)^2 + (2.576)^2 \\cdot (180)^2}')
doc.add_paragraph('n = \\frac{2687489280}{2812275 + 214999.14} = 887.76')

ans_c = doc.add_paragraph()
ans_c.add_run('Respuesta: ').bold = True
ans_c.add_run('Redondeando al entero superior, el tamaño de muestra es de 888 encuestas.')

# Tarea D
q_d = doc.add_paragraph()
q_d.add_run('d) Si el presupuesto máximo asignado para este estudio es de $15,000, ¿es viable realizar esta investigación con los parámetros exigidos? Justifique su respuesta con el cálculo del costo.').bold = True

doc.add_paragraph().add_run('Solución d):').bold = True
doc.add_paragraph('Costo Proyectado = 888 encuestas * $25.00 = $22,200.00')

ans_d = doc.add_paragraph()
ans_d.add_run('Respuesta: ').bold = True
ans_d.add_run('La investigación no es viable. El proyecto requiere una inversión de $22,200.00, superando el límite del presupuesto de $15,000.00 por una diferencia de $7,200.00. El Ministerio tendría que flexibilizar los parámetros.')

# Tarea 3
h2_3 = doc.add_heading('Tarea 3: Análisis de Sensibilidad y Criterio Económico', level=2)
for run in h2_3.runs: run.font.name = 'Times New Roman'; run.font.color.rgb = RGBColor(0,0,0); run.font.size = Pt(12)

# Tarea E
q_e = doc.add_paragraph()
q_e.add_run('e) Si se decide reducir el Nivel de Confianza de la Tarea 2 del 99% al 95%, ¿cuántas encuestas se ahorraría el Ministerio y cuánto dinero representaría ese ahorro?').bold = True

doc.add_paragraph().add_run('Solución e):').bold = True
doc.add_paragraph('Para un 95% de confianza, el nuevo Z = 1.96. Recalculamos el tamaño de muestra para la Tarea 2:')

inst = doc.add_paragraph('Fórmulas base para Word:')
inst.runs[0].font.color.rgb = RGBColor(100, 100, 100); inst.runs[0].italic = True

doc.add_paragraph('n = \\frac{12500 \\cdot (1.96)^2 \\cdot (180)^2}{12499 \\cdot (15)^2 + (1.96)^2 \\cdot (180)^2}')
doc.add_paragraph('n = \\frac{1555848000}{2812275 + 124467.84} = 529.78')

p_e = doc.add_paragraph('El nuevo requerimiento de muestra bajaría a 530 encuestas.\n')
p_e.add_run('Ahorro en encuestas: ').bold = True
p_e.add_run('888 (original) - 530 (nueva) = 358 encuestas ahorradas.\n')
p_e.add_run('Ahorro monetario: ').bold = True
p_e.add_run('358 encuestas * $25.00 = $8,950.00 de ahorro.')

# Tarea F
q_f = doc.add_paragraph()
q_f.add_run('f) Si no se tuviera el dato del 35% de acceso al crédito en la Tarea 1 y tuviera que asumir máxima variabilidad, ¿en qué porcentaje aumentaría el tamaño de la muestra respecto al cálculo original?').bold = True

doc.add_paragraph().add_run('Solución f):').bold = True
doc.add_paragraph('Asumir "máxima variabilidad" implica que p = 0.5 y q = 0.5. Recalculamos la Tarea 1:')

inst = doc.add_paragraph('Fórmulas base para Word:')
inst.runs[0].font.color.rgb = RGBColor(100, 100, 100); inst.runs[0].italic = True

doc.add_paragraph('n = \\frac{12500 \\cdot (1.96)^2 \\cdot 0.5 \\cdot 0.5}{12499 \\cdot (0.04)^2 + (1.96)^2 \\cdot 0.5 \\cdot 0.5}')
doc.add_paragraph('n = \\frac{12005}{19.9984 + 0.9604} = 572.79')

doc.add_paragraph('El requerimiento muestral incrementaría a 573 encuestas.\nIncremento absoluto: 573 - 524 = 49 encuestas extra.')

doc.add_paragraph('Variación porcentual = (49 / 524) * 100% = 9.35%')

ans_f = doc.add_paragraph()
ans_f.add_run('Respuesta: ').bold = True
ans_f.add_run('El costo estadístico en el tamaño de muestra sufriría un aumento de un 9.35% extra respecto al diseño inicial, debido a la incertidumbre total.')

doc.save('Solucion_Tarea_Muestreo_Word.docx')
print("¡Archivo DOCX creado exitosamente!")
