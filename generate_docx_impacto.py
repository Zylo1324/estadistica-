import os
import subprocess
import sys

# Instalar python-docx si no existe
try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Configurar márgenes
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Configurar fuente general
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11.5)

# Título
h1 = doc.add_heading('Actividad: Evaluación del Impacto Económico Post-Crisis', level=1)
h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in h1.runs:
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.size = Pt(14)

doc.add_paragraph()

# Caso
caso = doc.add_paragraph()
caso.add_run('Contexto: ').bold = True
caso.add_run('Tras una fuerte fluctuación en los precios internacionales de las materias primas, el Ministerio de Economía y Finanzas (MEF) necesita diseñar un paquete de transferencias monetarias para mitigar el impacto en las familias vulnerables. El presupuesto es limitado y el tiempo apremia. \n\nComo economistas del equipo técnico, se les pide determinar el tamaño de muestra necesario para dos estudios críticos en una región de 500,000 hogares:')

# Escenario A
h2_a = doc.add_heading('Escenario A: Estimación de la Tasa de Pobreza (Proporción)', level=2)
for run in h2_a.runs: run.font.name = 'Times New Roman'; run.font.color.rgb = RGBColor(0,0,0); run.font.size = Pt(12)

p_info = doc.add_paragraph('El gobierno necesita saber qué porcentaje de hogares ha caído por debajo de la línea de pobreza este trimestre para asignar el presupuesto total del subsidio.\n')
p_info.add_run('Información disponible: ').bold = True
p_info.add_run('En el censo anterior, la tasa de pobreza en esa región era del 25% (p = 0.25).\n')
p_info.add_run('Requerimiento: ').bold = True
p_info.add_run('El Ministro exige un nivel de confianza del 95% y un margen de error máximo del 3% (0.03).')

# Escenario B
h2_b = doc.add_heading('Escenario B: Estimación del Gasto Promedio en Alimentos (Media)', level=2)
for run in h2_b.runs: run.font.name = 'Times New Roman'; run.font.color.rgb = RGBColor(0,0,0); run.font.size = Pt(12)

p_info2 = doc.add_paragraph('Para ajustar el monto exacto del bono alimentario, se necesita estimar el gasto promedio mensual en la canasta básica.\n')
p_info2.add_run('Información disponible: ').bold = True
p_info2.add_run('Un estudio piloto pequeño sugiere que la desviación estándar del gasto en alimentos es de $120.\n')
p_info2.add_run('Requerimiento: ').bold = True
p_info2.add_run('Se solicita un nivel de confianza del 95% y un error máximo admisible de solo $10.')

# Pregunta 1
doc.add_heading('1. Tamaño de muestra sin conocer el tamaño de la población (N desconocido)', level=3)

q1 = doc.add_paragraph()
q1.add_run('¿A cuántos hogares debemos encuestar para que el subsidio no se quede corto ni se desperdicie? (calcule el tamaño de muestra sin conocer el tamaño de la población en ambos escenarios)').bold = True

doc.add_paragraph().add_run('Solución Escenario A (Proporción):').bold = True
doc.add_paragraph('Proporción esperada (p): 0.25\nProbabilidad de fracaso (q): 0.75\nConfianza (NC): 95%, cuyo valor Z es 1.96\nError (E): 0.03')
doc.add_paragraph('Fórmulas base para Word (Cópialas y presiona Alt + = para convertirlas automáticamente):').runs[0].font.color.rgb = RGBColor(100, 100, 100)
doc.add_paragraph('n = \\frac{Z^2 \\cdot p \\cdot q}{E^2}')
doc.add_paragraph('n = \\frac{(1.96)^2 \\cdot 0.25 \\cdot 0.75}{(0.03)^2} = \\frac{3.8416 \\cdot 0.1875}{0.0009} = \\frac{0.7203}{0.0009} = 800.33')
ans1a = doc.add_paragraph()
ans1a.add_run('Respuesta (A): ').bold = True
ans1a.add_run('Redondeando al entero superior, se requiere encuestar a 801 hogares.')

doc.add_paragraph().add_run('Solución Escenario B (Media):').bold = True
doc.add_paragraph('Desviación Estándar (\u03C3): $120\nConfianza (NC): 95%, cuyo valor Z es 1.96\nError (E): 10')
doc.add_paragraph('Fórmulas base para Word:').runs[0].font.color.rgb = RGBColor(100, 100, 100)
doc.add_paragraph('n = \\frac{Z^2 \\cdot \\sigma^2}{E^2}')
doc.add_paragraph('n = \\frac{(1.96)^2 \\cdot (120)^2}{(10)^2} = \\frac{3.8416 \\cdot 14400}{100} = \\frac{55319.04}{100} = 553.19')
ans1b = doc.add_paragraph()
ans1b.add_run('Respuesta (B): ').bold = True
ans1b.add_run('Redondeando al entero superior, se requiere encuestar a 554 hogares.')

# Pregunta 2
doc.add_heading('2. Tamaño de muestra asumiendo sin datos previos (p = 0.5) - Solo Escenario A', level=3)

q2 = doc.add_paragraph()
q2.add_run('Calcula nuevamente el tamaño de muestra asumiendo que no tienen datos previos (p = 0.5).').bold = True

doc.add_paragraph().add_run('Solución:').bold = True
doc.add_paragraph('Si no hay datos previos, asumimos máxima varianza, donde p = 0.5 y q = 0.5.')
doc.add_paragraph('Fórmulas base para Word:').runs[0].font.color.rgb = RGBColor(100, 100, 100)
doc.add_paragraph('n = \\frac{(1.96)^2 \\cdot 0.5 \\cdot 0.5}{(0.03)^2} = \\frac{3.8416 \\cdot 0.25}{0.0009} = \\frac{0.9604}{0.0009} = 1067.11')
ans2 = doc.add_paragraph()
ans2.add_run('Respuesta: ').bold = True
ans2.add_run('Si no tuviéramos datos previos, el tamaño de la muestra subiría a 1068 hogares.')

# Pregunta 3
doc.add_heading('3. Tamaño de muestra al 99% de confianza (N desconocido)', level=3)

q3 = doc.add_paragraph()
q3.add_run('Calcular a hora también al 99% de confianza en ambos escenarios.').bold = True

doc.add_paragraph('Para un Nivel de Confianza del 99%, el valor estadístico Z es de 2.576.')

doc.add_paragraph().add_run('Solución Escenario A (Proporción):').bold = True
doc.add_paragraph('Fórmulas base para Word:').runs[0].font.color.rgb = RGBColor(100, 100, 100)
doc.add_paragraph('n = \\frac{(2.576)^2 \\cdot 0.25 \\cdot 0.75}{(0.03)^2} = \\frac{6.635776 \\cdot 0.1875}{0.0009} = \\frac{1.244208}{0.0009} = 1382.45')
ans3a = doc.add_paragraph()
ans3a.add_run('Respuesta (A): ').bold = True
ans3a.add_run('Se requiere encuestar a 1383 hogares para el nivel exigido del 99%.')

doc.add_paragraph().add_run('Solución Escenario B (Media):').bold = True
doc.add_paragraph('Fórmulas base para Word:').runs[0].font.color.rgb = RGBColor(100, 100, 100)
doc.add_paragraph('n = \\frac{(2.576)^2 \\cdot (120)^2}{(10)^2} = \\frac{6.635776 \\cdot 14400}{100} = \\frac{95555.1744}{100} = 955.55')
ans3b = doc.add_paragraph()
ans3b.add_run('Respuesta (B): ').bold = True
ans3b.add_run('Se requiere encuestar a 956 hogares para el nivel exigido del 99%.')

# Pregunta 4
doc.add_heading('4. Población conocida (N = 500,000) - 95% Confianza', level=3)

q4 = doc.add_paragraph()
q4.add_run('Dado que la población es de 500,000, evaluar si es necesario aplicar la fórmula para el tamaño de la población conocida (en ambos casos).').bold = True

doc.add_paragraph().add_run('Solución Escenario A:').bold = True
doc.add_paragraph('Fórmulas base para Word:').runs[0].font.color.rgb = RGBColor(100, 100, 100)
doc.add_paragraph('n = \\frac{N \\cdot Z^2 \\cdot p \\cdot q}{(N-1) \\cdot E^2 + Z^2 \\cdot p \\cdot q}')
doc.add_paragraph('n = \\frac{500000 \\cdot (1.96)^2 \\cdot 0.25 \\cdot 0.75}{(499999) \\cdot (0.03)^2 + (1.96)^2 \\cdot 0.25 \\cdot 0.75}')
doc.add_paragraph('n = \\frac{360150}{449.9991 + 0.7203} = \\frac{360150}{450.7194} = 799.05')

doc.add_paragraph().add_run('Solución Escenario B:').bold = True
doc.add_paragraph('Fórmulas base para Word:').runs[0].font.color.rgb = RGBColor(100, 100, 100)
doc.add_paragraph('n = \\frac{N \\cdot Z^2 \\cdot \\sigma^2}{(N-1) \\cdot E^2 + Z^2 \\cdot \\sigma^2}')
doc.add_paragraph('n = \\frac{500000 \\cdot (1.96)^2 \\cdot (120)^2}{(499999) \\cdot (10)^2 + (1.96)^2 \\cdot (120)^2}')
doc.add_paragraph('n = \\frac{27659520000}{49999900 + 55319.04} = \\frac{27659520000}{50055219.04} = 552.58')

ans4 = doc.add_paragraph()
ans4.add_run('Evaluación: ').bold = True
ans4.add_run('Los resultados ajustados (800 y 553 respectivamente) son prácticamente idénticos a los de la población infinita (801 y 554). Dado que la población (N = 500,000) es muy extensa, el tamaño de la muestra original (n) representa apenas el 0.16% y 0.11% de la población total (ampliamente menor al umbral de sensibilidad del 5%). Por consiguiente, matemáticamente NO es estrictamente necesario aplicar la fórmula para la población finita. La fórmula de población infinita es más que suficiente en ambos casos.')

# Pregunta 5 y 6
doc.add_heading('Análisis y Criterio Económico', level=2)

q5 = doc.add_paragraph()
q5.add_run('¿Estarían dispuestos a arriesgar la precisión política por ahorrar costos de encuesta?').bold = True

doc.add_paragraph('Desde un punto de vista técnico-económico, se trata de balancear el costo contra el riesgo. Reducir la precisión disminuiría el tamaño de la muestra (por ejemplo al reducir la confianza del 99% al 95%) y con ello bajarían los costos investigativos. Sin embargo, en el diseño de un paquete de transferencias monetarias de emergencia, dar un subsidio deficiente o cometer un error y dejar familias vulnerables excluidas posee un costo social y político altísimo. Dado que el grado de sensibilidad del gasto público es crítico, el beneficio social de hacerlo correctamente a un nivel del 95% o superior sin duda compensa y justifica el aumento que supone aplicar las encuestas extra.')

q6 = doc.add_paragraph()
q6.add_run('Explique por qué la información previa tiene valor económico.').bold = True

doc.add_paragraph('Tener información previa disminuye asombrosamente la incertidumbre de la varianza. En el Escenario A, al conocer la tasa de pobreza anterior (p = 0.25), requirimos levantar solo 801 encuestas. Por contra, si no se tuviera el dato, nos vemos obligados a asumir la máxima incertidumbre estadística posible (p = 0.5) lo cual aumenta el requerimiento a 1068 encuestas. Esa brecha de 267 elementos adicionales requiere obligatoriamente una inversión económica extra en sueldos de encuestadores, viáticos y logística por parte del estado. Aquello manifiesta un principio vital de la econometría: la información y caracterización estadística de buena calidad "ahorra encuestas" directamente, dotando a la data de un evidente valor económico directo en ahorro de gestión.')

doc.save('Solucion_Evaluacion_Impacto_Economico.docx')
print("¡Archivo DOCX creado exitosamente!")
